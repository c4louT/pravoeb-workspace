#!/usr/bin/env python3
"""
Phase 5k Stage 1: Extract clause corpus from 69 .docx templates.

Parses each .docx in contracts/templates_docx/, identifies section headers
(ALL-CAPS numbered lines like "1.ПРЕДМЕТ ДОГОВОРА"), splits into clauses
(numbered like 1.1, 1.2, 2.1.1), tags each with metadata from templates_index.json,
and emits contracts/clauses_corpus.json.

Usage:
    python3 tools/extract_corpus.py
"""
from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path

try:
    import docx  # python-docx
except ImportError:
    raise SystemExit("pip install python-docx")

ROOT = Path(__file__).resolve().parents[1]
DOCX_DIR = ROOT / "contracts" / "templates_docx"
INDEX_FILE = ROOT / "contracts" / "templates_index.json"
OUT_FILE = ROOT / "contracts" / "clauses_corpus.json"

# Numbered section header: "1.ПРЕДМЕТ ДОГОВОРА" or "1. Предмет Договора".
SECTION_RE = re.compile(
    r"^\s*(\d{1,2})\s*\.\s*([А-ЯЁA-Z][А-ЯЁA-Za-zа-яё\s\-,№()/\"«»]{2,80}?)\.?\s*$"
)
# Bare ALL-CAPS section header without leading number: "ПРЕДМЕТ ДОГОВОРА"
BARE_SECTION_RE = re.compile(
    r"^\s*([А-ЯЁA-Z][А-ЯЁA-Z\s\-,№()/\"«»]{2,80}?)\.?\s*$"
)
CLAUSE_RE = re.compile(r"^\s*(\d{1,2}(?:\.\d{1,3}){1,3})\.?\s*(.*)$", re.DOTALL)
PLACEHOLDER_RE = re.compile(r"\{[a-zA-Z_][a-zA-Z0-9_]*\}")

# Canonical section names. Keys are lowercase normalized titles; values = canonical slug + priority order.
SECTION_MAP = {
    "предмет договора": ("subject", 1),
    "предмет договора услуг": ("subject", 1),
    "предмет": ("subject", 1),
    "срок договора": ("term", 2),
    "срок": ("term", 2),
    "сроки": ("term", 2),
    "сроки оказания услуг": ("term", 2),
    "права и обязанности сторон": ("rights_duties", 3),
    "права и обязанности": ("rights_duties", 3),
    "обязанности исполнителя": ("rights_duties", 3),
    "обязанности сторон": ("rights_duties", 3),
    "порядок оказания услуг": ("rights_duties", 3),
    "порядок сдачи-приемки услуг": ("acceptance", 4),
    "сдача-приемка услуг": ("acceptance", 4),
    "приемка": ("acceptance", 4),
    "стоимость услуг и порядок расчетов": ("payment", 5),
    "стоимость и порядок расчетов": ("payment", 5),
    "стоимость и порядок оплаты": ("payment", 5),
    "вознаграждение": ("payment", 5),
    "цена договора и порядок расчетов": ("payment", 5),
    "цена": ("payment", 5),
    "размер и порядок выплаты вознаграждения": ("payment", 5),
    "отчуждение исключительного права": ("rights_transfer", 6),
    "передача прав": ("rights_transfer", 6),
    "передача исключительных прав": ("rights_transfer", 6),
    "исключительные права": ("rights_transfer", 6),
    "исключительное право": ("rights_transfer", 6),
    "права на результат": ("rights_transfer", 6),
    "смежные права": ("rights_transfer", 6),
    "права на рид": ("rights_transfer", 6),
    "персональные данные": ("personal_data", 7),
    "персональные данные. личные неимущественные права": ("personal_data", 7),
    "персональные данные личные неимущественные права": ("personal_data", 7),
    "ответственность сторон": ("liability", 8),
    "ответственность": ("liability", 8),
    "обстоятельства непреодолимой силы": ("force_majeure", 9),
    "форс-мажор": ("force_majeure", 9),
    "конфиденциальность": ("confidentiality", 10),
    "антикоррупционная оговорка": ("anticorruption", 11),
    "антикоррупционные условия": ("anticorruption", 11),
    "расторжение договора": ("termination", 12),
    "расторжение": ("termination", 12),
    "порядок расторжения договора": ("termination", 12),
    "прочие условия": ("misc", 13),
    "заключительные положения": ("misc", 13),
    "заключительные условия": ("misc", 13),
    "разрешение споров": ("disputes", 14),
    "порядок разрешения споров": ("disputes", 14),
    "уведомления": ("notices", 15),
    "адреса, реквизиты и подписи сторон": ("requisites", 99),
    "реквизиты и подписи сторон": ("requisites", 99),
    "реквизиты сторон": ("requisites", 99),
    "подписи сторон": ("requisites", 99),
    # Additional common variants
    "действие договора": ("term", 2),
    "срок действия договора": ("term", 2),
    "условия изменения и расторжения договора": ("termination", 12),
    "условия конфиденциальности": ("confidentiality", 10),
    "положения о конфиденциальности и защите": ("confidentiality", 10),
    "финансовые условия": ("payment", 5),
    "финансовые условия и порядок расчетов": ("payment", 5),
    "стоимость работы и порядок расчетов": ("payment", 5),
    "расчеты и порядок оплаты": ("payment", 5),
    "размер и порядок оплаты работ": ("payment", 5),
    "стоимость договора и порядок расчета": ("payment", 5),
    "дополнительные условия и гарантии": ("guarantees", 11),
    "гарантии": ("guarantees", 11),
    "гарантии сторон": ("guarantees", 11),
    "общие условия договора": ("misc", 13),
    "общие положения": ("misc", 13),
    "особые условия": ("misc", 13),
    "дополнительные условия": ("misc", 13),
    "обязательства сторон": ("rights_duties", 3),
    "обязанности заказчика": ("rights_duties", 3),
    "права и обязательства сторон": ("rights_duties", 3),
    "определение терминов": ("definitions", 0),
    "порядок рассмотрения споров": ("disputes", 14),
    "предоставляемая лицензия": ("rights_transfer", 6),
    "произведения": ("subject", 1),
}


def normalize_section_title(title: str) -> str:
    t = title.lower().strip()
    t = re.sub(r"\s+", " ", t)
    t = t.rstrip(".,:; ")
    return t


def classify_section(title: str) -> tuple[str, int]:
    norm = normalize_section_title(title)
    if norm in SECTION_MAP:
        return SECTION_MAP[norm]
    # Fuzzy match by substring
    for key, val in SECTION_MAP.items():
        if key in norm or norm in key:
            return val
    return ("other", 50)


def clean_text(t: str) -> str:
    # Remove zero-width spaces and normalize non-breaking spaces
    t = t.replace("\u200b", "").replace("\xa0", " ")
    return t


def iter_paragraphs(doc: "docx.document.Document"):
    # Body paragraphs
    for p in doc.paragraphs:
        t = clean_text(p.text)
        for line in t.split("\n"):
            line = line.strip()
            if line:
                yield line
    # Table cell paragraphs (many templates embed clauses inside single-row tables)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = clean_text(p.text)
                    for line in t.split("\n"):
                        line = line.strip()
                        if line:
                            yield line


def split_on_internal_numbering(text: str) -> list[str]:
    """If a paragraph contains multiple numbered sub-clauses (e.g. '2.1.1. foo 2.1.2. bar'),
    or inline section headers ('1.ПРЕДМЕТ ДОГОВОРА    1.1. foo'),
    split them into separate lines."""
    # First split on inline section headers: "N.CAPS  " → newlines around it
    # Match: start or whitespace, then N., then 3-60 chars of Russian caps, then 2+ spaces or tab
    text2 = re.sub(
        r"(?<![\d.])(\d{1,2}\.[А-ЯЁ][А-ЯЁ\s\-,№()/\"«»]{2,80}?)(?=\s{2,}|\t|$)",
        r"\n\1\n",
        text,
    )
    # Then split on clause numbering like " 1.1." " 2.3.1."
    text3 = re.sub(
        r"(?<=[\.\s])(\d{1,2}(?:\.\d{1,3}){1,3}\.)\s*",
        r"\n\1 ",
        text2,
    )
    lines = [l.strip() for l in text3.split("\n") if l.strip()]
    return lines


def extract_clauses_from_docx(path: Path) -> list[dict]:
    doc = docx.Document(str(path))
    lines: list[str] = []
    for raw in iter_paragraphs(doc):
        for ln in split_on_internal_numbering(raw):
            lines.append(ln)

    clauses: list[dict] = []
    current_section_num: int | None = None
    current_section_title: str = ""
    current_section_slug: str = "other"
    current_section_order: int = 50
    current_clause: dict | None = None

    def flush():
        nonlocal current_clause
        if current_clause and current_clause["text"].strip():
            clauses.append(current_clause)
        current_clause = None

    auto_section_num = 0

    for ln in lines:
        sm = SECTION_RE.match(ln)
        if sm:
            flush()
            current_section_num = int(sm.group(1))
            current_section_title = sm.group(2).strip().rstrip(".").strip()
            current_section_slug, current_section_order = classify_section(current_section_title)
            auto_section_num = max(auto_section_num, current_section_num)
            continue

        # Bare ALL-CAPS section header without number
        bm = BARE_SECTION_RE.match(ln)
        if bm and len(ln) < 85:
            title = bm.group(1).strip().rstrip(".").strip()
            slug, order = classify_section(title)
            # Only accept if it maps to a known section (avoids false positives from ALL-CAPS prose)
            if slug != "other":
                flush()
                auto_section_num += 1
                current_section_num = auto_section_num
                current_section_title = title
                current_section_slug = slug
                current_section_order = order
                continue

        cm = CLAUSE_RE.match(ln)
        if cm and current_section_num is not None:
            # Only accept if clause number starts with the section number
            num = cm.group(1)
            first = int(num.split(".")[0])
            if first != current_section_num:
                # Belongs to current clause as continuation
                if current_clause:
                    current_clause["text"] += "\n" + ln
                continue
            flush()
            current_clause = {
                "clause_num": num,
                "text": cm.group(2).strip(),
                "section_num": current_section_num,
                "section_title": current_section_title,
                "section_slug": current_section_slug,
                "section_order": current_section_order,
            }
            continue

        # Plain continuation text
        if current_clause:
            current_clause["text"] += "\n" + ln
        elif current_section_num is not None:
            # Section intro prose without a clause number
            current_clause = {
                "clause_num": f"{current_section_num}.0",
                "text": ln,
                "section_num": current_section_num,
                "section_title": current_section_title,
                "section_slug": current_section_slug,
                "section_order": current_section_order,
            }

    flush()
    return clauses


def infer_contract_type(label: str, file_id: str) -> str:
    l = (label or "").lower()
    f = (file_id or "").lower()
    rules = [
        ("актёр", "actor"), ("актер", "actor"), ("акт" , "actor"),
        ("режиссёр", "director"), ("режиссер", "director"), ("director", "director"),
        ("оператор", "dop"), ("dop", "dop"), ("камер", "dop"),
        ("звук", "sound"), ("sound", "sound"),
        ("композитор", "composer"), ("composer", "composer"),
        ("монтаж", "editor"), ("editor", "editor"),
        ("сценар", "screenplay"), ("script", "screenplay"), ("litpro", "screenplay"),
        ("сценограф", "production_designer"), ("художник", "art"),
        ("костюм", "costume"), ("грим", "makeup"),
        ("свет", "light"), ("осветит", "light"),
        ("продюсер", "producer"), ("администратор", "administrator"),
        ("хореограф", "choreographer"),
        ("лицензи", "license"),
        ("согласие", "consent"),
        ("backstage", "backstage"),
        ("отчужден", "alienation"),
    ]
    for needle, slug in rules:
        if needle in l or needle in f:
            return slug
    return "other"


def main():
    idx = json.load(INDEX_FILE.open(encoding="utf-8"))
    tmeta = {t["id"]: t for t in idx["templates"]}

    all_clauses: list[dict] = []
    per_template_counts: dict[str, int] = {}
    section_counts: dict[str, int] = defaultdict(int)
    warnings: list[str] = []

    for t in idx["templates"]:
        fid = t["id"]
        fname = t.get("file") or f"{fid}.docx"
        path = DOCX_DIR / fname
        if not path.exists():
            warnings.append(f"Missing .docx: {fname}")
            continue

        try:
            clauses = extract_clauses_from_docx(path)
        except Exception as exc:
            warnings.append(f"Failed to parse {fname}: {exc}")
            continue

        tax_form = t.get("tax_form", "").strip() or "UNK"
        role_group = t.get("role_group", "").strip() or "UNK"
        contract_type = infer_contract_type(t.get("label", ""), fid)
        label = t.get("label", fid)

        for c in clauses:
            placeholders = PLACEHOLDER_RE.findall(c["text"])
            cid = f"{fid}__{c['clause_num'].replace('.', '_')}"
            all_clauses.append({
                "id": cid,
                "template_id": fid,
                "template_label": label,
                "template_file": fname,
                "contract_type": contract_type,
                "tax_form": tax_form,
                "role_group": role_group,
                "section_num": c["section_num"],
                "section_title": c["section_title"],
                "section_slug": c["section_slug"],
                "section_order": c["section_order"],
                "clause_num": c["clause_num"],
                "text": c["text"].strip(),
                "placeholders": sorted(set(placeholders)),
                "char_count": len(c["text"]),
            })
            section_counts[c["section_slug"]] += 1
        per_template_counts[fid] = len(clauses)

    corpus = {
        "version": "5k-v1",
        "generated_by": "tools/extract_corpus.py",
        "total_templates": len(per_template_counts),
        "total_clauses": len(all_clauses),
        "section_counts": dict(section_counts),
        "warnings": warnings,
        "clauses": all_clauses,
    }

    OUT_FILE.write_text(
        json.dumps(corpus, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print(f"✓ Wrote {OUT_FILE.relative_to(ROOT)}")
    print(f"  Templates parsed : {len(per_template_counts)} / {len(idx['templates'])}")
    print(f"  Total clauses    : {len(all_clauses)}")
    print(f"  Section counts   : {dict(section_counts)}")
    if warnings:
        print(f"  Warnings ({len(warnings)}):")
        for w in warnings[:10]:
            print(f"    - {w}")
        if len(warnings) > 10:
            print(f"    ... +{len(warnings) - 10} more")

    # Per-template size distribution
    sizes = sorted(per_template_counts.values())
    if sizes:
        print(f"  Clauses/template : min={sizes[0]} median={sizes[len(sizes)//2]} max={sizes[-1]}")


if __name__ == "__main__":
    main()
