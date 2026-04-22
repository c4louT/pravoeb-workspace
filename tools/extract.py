#!/usr/bin/env python3
"""
Универсальный извлекатель текста из договоров.
Использование:
    ./extract.py <путь-к-файлу> [--max-chars N]

Поддерживает:
  - .pdf   (через pypdf)
  - .docx  (через python-docx)
  - .txt / .md / прочее — читает как UTF-8
"""
import sys
import argparse
import os
from pathlib import Path


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            text = f"[page {i} extract error: {e}]"
        parts.append(f"\n--- Page {i} ---\n{text}")
    return "\n".join(parts)


def extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    lines = []

    # Параграфы
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)

    # Таблицы
    for t_idx, table in enumerate(doc.tables, 1):
        lines.append(f"\n[Таблица {t_idx}]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("file", help="Путь к файлу")
    ap.add_argument("--max-chars", type=int, default=0,
                    help="Максимум символов (0 = без лимита)")
    args = ap.parse_args()

    path = Path(args.file).expanduser().resolve()
    if not path.exists():
        sys.exit(f"Файл не найден: {path}")

    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            text = extract_pdf(path)
        elif ext == ".docx":
            text = extract_docx(path)
        elif ext in {".txt", ".md", ".rtf", ""}:
            text = extract_text(path)
        else:
            # Попробуем как текст
            text = extract_text(path)
    except Exception as e:
        sys.exit(f"Ошибка извлечения: {e}")

    if args.max_chars and len(text) > args.max_chars:
        text = text[:args.max_chars] + f"\n\n[...обрезано, всего {len(text)} символов]"

    print(f"# Файл: {path.name}")
    print(f"# Размер: {path.stat().st_size} байт, {len(text)} символов")
    print(f"# Тип: {ext or 'text'}")
    print()
    print(text)


if __name__ == "__main__":
    main()
