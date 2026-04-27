#!/usr/bin/env python3
"""Extract text from all .docx/.docm/.doc/.pdf files into per-file .txt for analysis.
Output: /tmp/phase5b/text/<unique_base>.txt (one per unique base, prefer .docx/.docm > .doc > .pdf)"""
import os, re, subprocess, unicodedata
from pathlib import Path
from collections import defaultdict

SRC = Path("/tmp/phase5b/Договора")
DST = Path("/tmp/phase5b/text")
DST.mkdir(exist_ok=True, parents=True)

def normalize(name):
    n = unicodedata.normalize("NFC", name)
    for _ in range(5):
        n = re.sub(r"\.(pdf|docx|doc|docm|rtf)\b", " ", n, flags=re.I)
        n = re.sub(r"\s*\(\d+\)\s*", " ", n)
    n = re.sub(r"\s+", " ", n).strip()
    return n

def safe_filename(name):
    """Slug-ify Cyrillic name into ASCII-safe filename."""
    n = unicodedata.normalize("NFC", name)
    n = re.sub(r"[^\w\s-]", "", n, flags=re.UNICODE)
    n = re.sub(r"\s+", "_", n)
    return n[:200]

# Group files by unique base
groups = defaultdict(list)
for p in sorted(SRC.iterdir()):
    if p.is_file():
        base = normalize(p.name)
        groups[base].append(p)

# Pick best version per group: docx > docm > doc > pdf
def quality(p):
    e = p.suffix.lower()
    return {".docx": 4, ".docm": 3, ".doc": 2, ".pdf": 1, ".rtf": 0}.get(e, -1)

extracted = 0
errors = []
for base, paths in groups.items():
    paths.sort(key=quality, reverse=True)
    src = paths[0]
    out = DST / (safe_filename(base) + ".txt")
    if out.exists():
        extracted += 1
        continue
    text = ""
    ext = src.suffix.lower()
    try:
        if ext == ".docx":
            import docx
            d = docx.Document(str(src))
            text = "\n".join(p.text for p in d.paragraphs)
        elif ext == ".docm":
            import docx
            d = docx.Document(str(src))
            text = "\n".join(p.text for p in d.paragraphs)
        elif ext == ".doc":
            r = subprocess.run(["catdoc", "-w", str(src)], capture_output=True, text=True, timeout=30)
            text = r.stdout
            if not text.strip():
                r = subprocess.run(["antiword", str(src)], capture_output=True, text=True, timeout=30)
                text = r.stdout
        elif ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(str(src)) as pdf:
                text = "\n".join((page.extract_text() or "") for page in pdf.pages)
        else:
            errors.append((src.name, f"unsupported ext {ext}"))
            continue
        if text.strip():
            # NFC normalize
            text = unicodedata.normalize("NFC", text)
            out.write_text(text, encoding="utf-8")
            extracted += 1
        else:
            errors.append((src.name, "empty extraction"))
    except Exception as e:
        errors.append((src.name, str(e)[:200]))

print(f"Extracted: {extracted} / {len(groups)} unique bases")
print(f"Errors: {len(errors)}")
for n, e in errors[:30]:
    print(f"  {n}: {e}")
